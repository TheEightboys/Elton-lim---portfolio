from docx import Document
import os

doc = Document('Elton Lim (1).docx')

# Create output file
with open('docx_content.txt', 'w', encoding='utf-8') as f:
    f.write('=== TEXT CONTENT ===\n\n')
    for para in doc.paragraphs:
        if para.text.strip():
            f.write(para.text + '\n')
    
    f.write('\n\n=== TABLE CONTENT ===\n\n')
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            f.write(' | '.join(row_text) + '\n')
        f.write('\n')

print("Content extracted to docx_content.txt")
