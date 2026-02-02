from docx import Document
import os

doc = Document('Elton Lim (1).docx')

print('=== TEXT CONTENT ===')
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

# Also check tables
print('\n=== TABLE CONTENT ===')
for table in doc.tables:
    for row in table.rows:
        row_text = [cell.text for cell in row.cells]
        print(' | '.join(row_text))

# Extract images
os.makedirs('extracted_images', exist_ok=True)
count = 0
for rel in doc.part.rels.values():
    if 'image' in rel.reltype:
        count += 1
        ext = rel.target_ref.split('.')[-1]
        image_data = rel.target_part.blob
        with open(f'extracted_images/image_{count}.{ext}', 'wb') as f:
            f.write(image_data)
        print(f'Extracted: image_{count}.{ext}')

print(f'\nTotal images extracted: {count}')
